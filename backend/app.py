"""
app.py
------
FastAPI web server providing REST endpoints for document PII redaction,
synthetic replacement diff generation, and ground-truth benchmark evaluation.
"""

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
import shutil
import os
from docx import Document

from redactor import redact_docx, redact_text, filter_and_resolve_overlaps, analyzer, SUPPORTED_ENTITIES
from fake_generator import get_fake_value
from evaluator import evaluate_redaction, generate_evaluation_markdown

app = FastAPI(
    title="Privora Engine API",
    version="1.0.0",
    description="Enterprise API for PII Detection, Synthetic Replacement, and Document Sanitization"
)

# Enable CORS for local Next.js frontend development
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)


@app.get("/")
def get_service_status():
    """Health check endpoint returning API operational status and supported PII categories."""
    return {
        "status": "online",
        "service": "Privora Engine",
        "version": "1.0.0",
        "supported_categories": [
            "Full Names (PERSON)",
            "Email Addresses (EMAIL_ADDRESS)",
            "Phone Numbers (PHONE_NUMBER)",
            "Company Names (COMPANY)",
            "Physical/Mailing Addresses (ADDRESS)",
            "Social Security Numbers (SSN)",
            "Credit Card Numbers (CREDIT_CARD)",
            "Dates of Birth (DATE_OF_BIRTH)",
            "IP Addresses (IP_ADDRESS)"
        ]
    }


@app.post("/redact-full")
async def process_document_redaction(file: UploadFile = File(...)):
    """
    Processes an uploaded .docx or .txt document, detects PII entities,
    substitutes detected instances with realistic synthetic fakes,
    saves the redacted document, and returns summary stats + verification diffs.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file was selected.")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in [".docx", ".txt"]:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file extension '{ext}'. Supported formats: .docx, .txt"
        )

    input_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(input_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    sample_diffs = []
    entity_counts = {
        "PERSON": 0,
        "EMAIL_ADDRESS": 0,
        "PHONE_NUMBER": 0,
        "COMPANY": 0,
        "ADDRESS": 0,
        "SSN": 0,
        "CREDIT_CARD": 0,
        "DATE_OF_BIRTH": 0,
        "IP_ADDRESS": 0
    }

    if ext == ".docx":
        output_path, counts = redact_docx(input_path, OUTPUT_DIR)
        for k, v in counts.items():
            if k in entity_counts:
                entity_counts[k] += v

        # Generate verification sample diff snippets
        doc = Document(input_path)
        full_text = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                full_text.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text and p.text.strip():
                            full_text.append(p.text)

        combined_text = "\n".join(full_text[:30])
        results = analyzer.analyze(text=combined_text, language="en", entities=SUPPORTED_ENTITIES)
        filtered = filter_and_resolve_overlaps(results, combined_text)

        seen_orig = set()
        for r in filtered:
            orig = combined_text[r.start:r.end].strip()
            etype = r.entity_type
            if etype == "LOCATION":
                etype = "ADDRESS"
            elif etype == "ORGANIZATION":
                etype = "COMPANY"

            if orig not in seen_orig and etype in entity_counts:
                seen_orig.add(orig)
                fake_val = get_fake_value(etype, orig)
                sample_diffs.append({
                    "category": etype,
                    "original": orig,
                    "redacted": fake_val
                })
                if len(sample_diffs) >= 15:
                    break

    else:  # .txt
        with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

        redacted_text, counts = redact_text(text)
        output_filename = "redacted_" + file.filename
        output_path = os.path.join(OUTPUT_DIR, output_filename)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(redacted_text)

        for k, v in counts.items():
            if k in entity_counts:
                entity_counts[k] += v

        results = analyzer.analyze(text=text[:5000], language="en", entities=SUPPORTED_ENTITIES)
        filtered = filter_and_resolve_overlaps(results, text[:5000])
        seen_orig = set()
        for r in filtered:
            orig = text[r.start:r.end].strip()
            etype = r.entity_type
            if etype == "LOCATION":
                etype = "ADDRESS"
            elif etype == "ORGANIZATION":
                etype = "COMPANY"
            if orig not in seen_orig and etype in entity_counts:
                seen_orig.add(orig)
                fake_val = get_fake_value(etype, orig)
                sample_diffs.append({
                    "category": etype,
                    "original": orig,
                    "redacted": fake_val
                })
                if len(sample_diffs) >= 15:
                    break

    total_detected = sum(entity_counts.values())

    return {
        "success": True,
        "filename": file.filename,
        "output_filename": os.path.basename(output_path),
        "total_detected": total_detected,
        "entity_counts": entity_counts,
        "sample_diffs": sample_diffs
    }


@app.post("/redact")
async def redact_file_legacy(file: UploadFile = File(...)):
    """Legacy file download endpoint returning stream response."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename missing.")

    input_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(input_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    if file.filename.endswith(".docx"):
        output_path, entity_counts = redact_docx(input_path, OUTPUT_DIR)
        return FileResponse(
            output_path,
            filename=os.path.basename(output_path),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"X-Entity-Counts": str(entity_counts)}
        )
    else:
        with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        redacted_text, entity_counts = redact_text(text)
        out_txt = os.path.join(OUTPUT_DIR, "redacted_" + file.filename)
        with open(out_txt, "w", encoding="utf-8") as f:
            f.write(redacted_text)
        return FileResponse(
            out_txt,
            filename=os.path.basename(out_txt),
            media_type="text/plain",
            headers={"X-Entity-Counts": str(entity_counts)}
        )


@app.get("/download/{filename}")
def download_file(filename: str):
    """Downloads a processed redacted file from the outputs directory."""
    path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="File not found.")
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if filename.endswith(".docx") else "text/plain"
    return FileResponse(path, filename=filename, media_type=media_type)


@app.get("/evaluate")
def execute_benchmark():
    """Runs evaluation benchmark against sample_ground_truth.json."""
    docx_file = os.path.join(UPLOAD_DIR, "Red_Herring_Prospectus.docx")
    gt_file = "sample_ground_truth.json"

    if not os.path.exists(docx_file):
        raise HTTPException(status_code=404, detail="Prospectus document not found in uploads directory.")
    if not os.path.exists(gt_file):
        raise HTTPException(status_code=404, detail="Ground-truth dataset not found.")

    res = evaluate_redaction(docx_file, gt_file)
    generate_evaluation_markdown(res, "evaluation_report.md")
    return JSONResponse(content=res)