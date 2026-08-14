"""
redactor.py
-----------
Main redaction module that reads source text/DOCX documents,
detects PII across 9 categories using Presidio + custom recognizers,
and replaces detected instances with consistent realistic fakes.
"""

from docx import Document
import os
from presidio_analyzer import AnalyzerEngine
from presidio_analyzer.nlp_engine import NlpEngineProvider
from custom_recognizers import register_all_custom_recognizers
from fake_generator import get_fake_value

# Read SPACY_MODEL environment variable, defaulting to lightweight deployment model en_core_web_sm
SPACY_MODEL = os.getenv("SPACY_MODEL", "en_core_web_sm")

def create_analyzer_engine(model_name: str = None) -> AnalyzerEngine:
    """Configures and returns a Presidio AnalyzerEngine using the target spaCy model."""
    target_model = model_name or SPACY_MODEL
    config = {
        "nlp_engine_name": "spacy",
        "models": [{"lang_code": "en", "model_name": target_model}]
    }
    provider = NlpEngineProvider(nlp_configuration=config)
    nlp_engine = provider.create_engine()
    engine = AnalyzerEngine(nlp_engine=nlp_engine)
    register_all_custom_recognizers(engine)
    return engine

# Initialize default Analyzer Engine with configured model
analyzer = create_analyzer_engine()

# All supported entities for the assignment
SUPPORTED_ENTITIES = [
    "PERSON",
    "EMAIL_ADDRESS",
    "PHONE_NUMBER",
    "COMPANY",
    "ORGANIZATION",
    "ADDRESS",
    "LOCATION",
    "SSN",
    "CREDIT_CARD",
    "DATE_OF_BIRTH",
    "IP_ADDRESS",
    "AADHAAR"
]

IGNORE_WORDS = {
    "Name", "Email", "Phone", "Aadhaar", "Address", "Company", "SSN", "DOB",
    "IP", "Website", "Tel", "Fax", "Date", "Section", "Table", "Page", "Type",
    "Particulars", "General", "Term", "Description", "Offer", "Fresh Issue",
    "Total", "Issuer", "Public", "Equity", "Shares", "Board", "Auditors"
}


import re
from presidio_analyzer import RecognizerResult

LABEL_PREFIX_REGEX = re.compile(
    r"^(?:(?:Registered Office|Corporate Office|Registered and Corporate Office|Address)|[A-Za-z0-9&\s.-]+(?:Limited|LLP|Inc\.?|Corporation|Pvt\.?\s*Ltd\.?))\s*:\s*",
    re.IGNORECASE
)


def cleanup_formatting(text: str) -> str:
    """
    Cleans up leftover formatting artifacts, double spaces, and space-punctuation issues.
    """
    if not text:
        return text
    # 1. Collapse multiple inline spaces into a single space
    text = re.sub(r"[ \t]{2,}", " ", text)
    # 2. Fix space before punctuation (e.g. "word , word" -> "word, word")
    text = re.sub(r"\s+([,.:;])", r"\1", text)
    # 3. Ensure a space after colons if followed immediately by alphanumeric or plus
    text = re.sub(r"(:)([A-Za-z0-9+])", r"\1 \2", text)
    # 4. Clean up trailing spaces on lines
    text = re.sub(r" ?\n ?", "\n", text)
    return text.strip()


from custom_recognizers import register_all_custom_recognizers, PROSPECTUS_COMPANIES

COMPANY_SUFFIX_REGEX = re.compile(
    r"\b(?:Private\s+Limited|Pvt\.?\s*Ltd\.?|Limited|Ltd\.?|LLP|Inc\.?|Corporation|Corp\.?|Bank|Trust|Holdings|Associates|Research|Partners|Capital|Services|Logistics|Infra|Distriparks|Industries|Motors|Securities|Management|Solutions|Technologies|Ventures|Group|Fund)\b",
    re.IGNORECASE
)

HEADER_NON_COMPANY_WORDS = {
    "SIZE", "ELIGIBILITY", "E-MAIL", "EMAIL", "TELEPHONE", "DETAILS", "OFFER", "PUBLIC",
    "FRESH", "ISSUE", "TYPE", "TOTAL", "WEBSITE", "REGISTERED", "OFFICE", "CORPORATE",
    "CONTACT", "PERSON", "LISTING", "RISKS", "FINANCE", "STATEMENT", "BOARD", "AUDITORS",
    "ISSUER", "EQUITY", "SHARES", "GENERAL", "PARTICULARS", "TERMS", "DESCRIPTION",
    "SUMMARY", "RISK", "TABLE", "PAGE", "CLAUSE", "NUMBER", "AMOUNT", "PRICE", "PERCENTAGE",
    "BID", "PERIOD", "OBJECTS", "PROMOTERS", "INFORMATION", "SHARE", "NAME", "DATE", "IP",
    "SSN", "DOB", "CIN", "DIN", "PAN", "LEI", "ISIN", "GST", "GIFT", "SEBI", "RBI", "MCA",
    "BSE", "NSE", "CDSL", "NSDL", "STT", "ITR", "TDS", "FII", "DII", "NRI", "QIB", "NIB",
    "RIB", "EMPLOYEE", "RESERVATION", "PORTION", "NET", "FRESH ISSUE", "OFFER FOR SALE",
    "TOTAL OFFER", "PRE-ISSUE", "POST-ISSUE", "FACE VALUE", "TICK SIZE", "SOCIAL SECURITY NUMBER",
    "CREDIT CARD NUMBER", "GATEWAY IP ADDRESS", "PERMANENT ACCOUNT NUMBER", "DETAILS OF THE OFFER TO PUBLIC",
    "RED HERRING PROSPECTUS", "REGISTRAR", "BOARD OF DIRECTORS", "SERVER IP ADDRESS", "BANKERS", "CFO",
    "KEY MANAGERIAL PERSONNEL"
}


def is_valid_company_entity(entity_str: str, score: float = 1.0) -> bool:
    """
    Strictly validates company entity matches to suppress false positives on ordinary English words,
    table header labels, and non-company text.
    """
    s_clean = entity_str.strip().strip(".,:;'\"()")
    s_upper = s_clean.upper()

    if not s_clean or s_upper in HEADER_NON_COMPANY_WORDS:
        return False

    words = s_clean.split()
    if len(words) == 1 and not COMPANY_SUFFIX_REGEX.search(s_clean):
        if any(c.strip().upper() == s_upper for c in PROSPECTUS_COMPANIES):
            return True
        return False

    for c in PROSPECTUS_COMPANIES:
        c_up = c.strip().upper()
        if c_up == s_upper or c_up in s_upper or s_upper in c_up:
            return True

    if COMPANY_SUFFIX_REGEX.search(s_clean):
        return True

    return False


def filter_and_resolve_overlaps(results, text):
    """
    Filters out noise/labels and resolves overlapping entity spans,
    enforcing container priority (ADDRESS containers override sub-tokens)
    and giving precedence to higher confidence and longer spans.
    """
    filtered = []
    for res in results:
        # If match is an ADDRESS starting with a label prefix (e.g. Registered Office:),
        # adjust start offset so the label prefix remains unredacted.
        if res.entity_type in ("ADDRESS", "LOCATION"):
            matched_str = text[res.start:res.end]
            m = LABEL_PREFIX_REGEX.match(matched_str)
            if m:
                new_start = res.start + m.end()
                if new_start < res.end:
                    res = RecognizerResult(
                        entity_type=res.entity_type,
                        start=new_start,
                        end=res.end,
                        score=res.score
                    )

        # For PHONE_NUMBER candidates, expand start backward if preceding text contains leading '+' or '+ '
        elif res.entity_type == "PHONE_NUMBER":
            prefix_check = text[max(0, res.start - 4):res.start]
            m_plus = re.search(r"\+\s*$", prefix_check)
            if m_plus:
                new_start = max(0, res.start - (len(prefix_check) - m_plus.start()))
                res = RecognizerResult(
                    entity_type=res.entity_type,
                    start=new_start,
                    end=res.end,
                    score=res.score
                )

        original = text[res.start:res.end].strip()

        # Filter out invalid COMPANY / ORGANIZATION false positives
        if res.entity_type in ("COMPANY", "ORGANIZATION"):
            if not is_valid_company_entity(original, res.score):
                continue

        if original in IGNORE_WORDS or original.upper() in HEADER_NON_COMPANY_WORDS:
            continue
        if len(original) < 2:
            continue
        # Avoid single-word PERSON detections if low score
        if res.entity_type == "PERSON" and len(original.split()) < 2 and res.score < 0.90:
            continue
        filtered.append(res)

    # Sort key order:
    # 1. ADDRESS / LOCATION entities get container precedence (type_priority = 10) over sub-tokens (5)
    # 2. Score descending
    # 3. Span length descending
    def sort_key(r):
        type_priority = 10 if r.entity_type in ("ADDRESS", "LOCATION") else 5
        return (type_priority, r.score, r.end - r.start)

    sorted_res = sorted(filtered, key=sort_key, reverse=True)

    kept = []
    for r in sorted_res:
        overlap = False
        for k in kept:
            # Overlap check: spans overlap if not completely disjoint
            if not (r.end <= k.start or r.start >= k.end):
                overlap = True
                break
        if not overlap:
            kept.append(r)

    # Return sorted by start position ascending for orderly processing
    return sorted(kept, key=lambda r: r.start)


def redact_text(text: str):
    """
    Analyzes text for PII entities, replaces them with realistic fakes,
    applies formatting cleanup, and returns (redacted_text, entity_counts).
    """
    if not text or not text.strip():
        return text, {}

    results = analyzer.analyze(
        text=text,
        language="en",
        entities=SUPPORTED_ENTITIES
    )

    filtered_results = filter_and_resolve_overlaps(results, text)

    entity_counts = {}
    for res in filtered_results:
        entity_type = res.entity_type
        # Normalize entity names
        if entity_type == "LOCATION":
            entity_type = "ADDRESS"
        elif entity_type == "ORGANIZATION":
            entity_type = "COMPANY"

        entity_counts[entity_type] = entity_counts.get(entity_type, 0) + 1

    # Replace from end of text to start to keep character offsets valid
    sorted_for_replacement = sorted(
        filtered_results,
        key=lambda x: x.start,
        reverse=True
    )

    redacted_text = text
    for res in sorted_for_replacement:
        original = text[res.start:res.end]

        entity_type = res.entity_type
        if entity_type == "LOCATION":
            entity_type = "ADDRESS"
        elif entity_type == "ORGANIZATION":
            entity_type = "COMPANY"

        replacement = get_fake_value(entity_type, original)
        redacted_text = (
            redacted_text[:res.start]
            + replacement
            + redacted_text[res.end:]
        )

    # Apply general formatting cleanup for double spaces & punctuation spacing
    cleaned_redacted_text = cleanup_formatting(redacted_text)

    return cleaned_redacted_text, entity_counts


from docx.oxml.ns import qn
from docx.shared import RGBColor


def get_shd_fill(element):
    """Retrieves w:fill attribute from a w:shd element if present."""
    if element is None:
        return None
    shd = element.find(qn("w:shd"))
    if shd is not None:
        return shd.get(qn("w:fill"))
    return None


def is_red_or_dark_bg(fill_str: str) -> bool:
    """
    Returns True if the hex fill color represents a red or dark-red background
    (e.g., C00000, D20000, E10000, 800000) or general dark background.
    """
    if not fill_str or fill_str.lower() in ("auto", "none", "clear", "ffffff"):
        return False
    fill_hex = fill_str.strip().lstrip("#").upper()
    if len(fill_hex) == 6:
        try:
            r = int(fill_hex[0:2], 16)
            g = int(fill_hex[2:4], 16)
            b = int(fill_hex[4:6], 16)
            # Red dominant check: r >= 100 and r > g * 1.2 and r > b * 1.2
            is_red = (r >= 100 and r > g * 1.2 and r > b * 1.2)
            luminance = (r * 299 + g * 587 + b * 114) / 1000
            return is_red or luminance < 130
        except ValueError:
            pass
    return False


def apply_paragraph_redaction_and_styling(paragraph, fill_context: str = None):
    """
    Redacts PII in a paragraph while preserving font styles (bold, italic, size, name)
    and explicitly setting font color to WHITE if the background is red/dark-red.
    """
    if not paragraph.text or not paragraph.text.strip():
        return {}

    pPr = paragraph._p.find(qn("w:pPr"))
    p_fill = get_shd_fill(pPr)
    effective_fill = fill_context or p_fill

    had_white_text = False
    for run in paragraph.runs:
        if run.font and run.font.color and run.font.color.rgb == RGBColor(255, 255, 255):
            had_white_text = True
            break

    should_be_white = is_red_or_dark_bg(effective_fill) or had_white_text

    is_bold = None
    is_italic = None
    font_name = None
    font_size = None
    if paragraph.runs:
        first_run = paragraph.runs[0]
        is_bold = first_run.bold
        is_italic = first_run.italic
        if first_run.font:
            font_name = first_run.font.name
            font_size = first_run.font.size

    redacted_text, counts = redact_text(paragraph.text)
    paragraph.text = redacted_text

    for run in paragraph.runs:
        if is_bold is not None:
            run.bold = is_bold
        if is_italic is not None:
            run.italic = is_italic
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if should_be_white:
            run.font.color.rgb = RGBColor(255, 255, 255)

    return counts


def redact_docx(path: str, output_dir: str = "outputs"):
    """
    Reads a .docx file, detects PII in paragraphs and tables,
    replaces PII with fakes, preserves bold/italic styles,
    ensures red-background header cells have WHITE text, and saves redacted document.
    """
    doc = Document(path)
    total_entity_counts = {}

    # 1. Process document paragraphs
    for paragraph in doc.paragraphs:
        counts = apply_paragraph_redaction_and_styling(paragraph)
        for k, v in counts.items():
            total_entity_counts[k] = total_entity_counts.get(k, 0) + v

    # 2. Process table cells
    for table in doc.tables:
        for row in table.rows:
            trPr = row._tr.find(qn("w:trPr"))
            tr_fill = get_shd_fill(trPr)
            for cell in row.cells:
                tcPr = cell._tc.find(qn("w:tcPr"))
                tc_fill = get_shd_fill(tcPr)
                cell_fill = tc_fill or tr_fill
                for paragraph in cell.paragraphs:
                    counts = apply_paragraph_redaction_and_styling(paragraph, fill_context=cell_fill)
                    for k, v in counts.items():
                        total_entity_counts[k] = total_entity_counts.get(k, 0) + v

    os.makedirs(output_dir, exist_ok=True)
    filename = os.path.basename(path)
    output_path = os.path.join(output_dir, "redacted_" + filename)
    doc.save(output_path)

    return output_path, total_entity_counts